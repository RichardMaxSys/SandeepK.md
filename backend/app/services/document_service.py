from docx import Document
from weasyprint import HTML
import io

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(content_text: str) -> io.BytesIO:
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for line in content_text.split('\n'):
        line = line.strip()
        if not line:
            continue

        p = doc.add_paragraph()
        if line.isupper() and len(line) < 50: # Assume header
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue
        else:
            p.add_run(line)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_pdf(content_text: str) -> io.BytesIO:
    # Premium Recruiter-Grade HTML Template
    sections = content_text.split('\n')
    styled_content = ""
    for line in sections:
        line = line.strip()
        if not line:
            continue
        if line.isupper() and len(line) < 60:
            styled_content += f"<h2 style='color: #1e3a8a; border-bottom: 2px solid #e5e7eb; margin-top: 24px; margin-bottom: 12px; font-weight: 700; font-size: 14pt; letter-spacing: 0.05em;'>{line}</h2>"
        elif line.startswith('•') or line.startswith('-'):
            styled_content += f"<li style='margin-bottom: 6px; margin-left: 20px; list-style-type: disc;'>{line.lstrip('•- ')}</li>"
        else:
            styled_content += f"<p style='margin-bottom: 8px; font-size: 11pt;'>{line}</p>"

    html_content = f"""
    <html>
        <head>
            <style>
                @page {{ margin: 0.75in; }}
                body {{ font-family: 'Inter', 'Helvetica', 'Arial', sans-serif; line-height: 1.6; color: #111827; }}
                h2 {{ text-transform: uppercase; }}
                li {{ font-size: 11pt; }}
            </style>
        </head>
        <body>
            {styled_content}
        </body>
    </html>
    """
    pdf_stream = io.BytesIO()
    HTML(string=html_content).write_pdf(pdf_stream)
    pdf_stream.seek(0)
    return pdf_stream
